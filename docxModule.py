import docx
from docx import Document
from docx.shared import Pt
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def title(title):
    doc = docx.Document('Document.docx')
    paragraph1 = doc.add_paragraph()
    style1 = doc.styles['Normal']
    font1 = style1.font
    font1.name = 'Times New Roman'
    font1.size = Pt(28)
    paragraph1.style = doc.styles['Normal']
    doc.add_paragraph(title)
    doc.save('Document.docx')
    return "1"

def newPara():
    doc = docx.Document('Document.docx')
    doc.add_paragraph(" \n ")
    doc.save("Document.docx")
    return "2"


def addPara(text):
    document = Document('Document.docx')
    paragraph = document.add_paragraph()
    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    paragraph.style = document.styles['Normal']
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 2
    document.add_paragraph(text)
    document.save('Document.docx')
    return "1"


def paraSpace():
    doc = docx.Document('Document.docx')
    doc.save("Document.docx")
    return "5"

